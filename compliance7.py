# ===============================
# 1. Imports
# ===============================
import gradio as gr
import pdfplumber
import json
import yaml
import textwrap
import os
import tempfile
import subprocess
import sys
import json
import re
import pandas as pd
from docx import Document
import io
from datetime import datetime
import traceback
import torch
import time
import threading
from typing import Tuple
import numpy as np
from pathlib import Path
#from langgraph.graph import StateGraph
from typing import Literal,Optional, List
from transformers import AutoModelForCausalLM, AutoTokenizer, pipeline
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.prompts import PromptTemplate
from pydantic import BaseModel, Field, ValidationError
from langchain_core.runnables import RunnableLambda
from peft import PeftModel
from sentence_transformers import SentenceTransformer
from concurrent.futures import ThreadPoolExecutor
from sklearn.metrics.pairwise import cosine_similarity
# Global model and tokenizer variables to avoid reloading
GLOBAL_MODEL = None
GLOBAL_TOKENIZER = None
# Global LLM wrapper to reduce model loading
_LLM_WRAPPER = None
from datetime import datetime
DEBUG_LOG_FILE = "/home/gpu/Documents/Amal/compliance_llama/llm_debug.txt"

def write_debug_log(message: str):
    timestamp = datetime.now().strftime("[%Y-%m-%d %H:%M:%S]")
    with open(DEBUG_LOG_FILE, "a", encoding="utf-8") as f:
        f.write(f"{timestamp} {message}\n")

write_debug_log("Test log entry - should appear in llm_debug.txt")

def get_llm_wrapper(model=None, tokenizer=None):
    """Get or create a shared LLM wrapper to avoid reloading the model."""
    global _LLM_WRAPPER
    if _LLM_WRAPPER is None:
        if model is None or tokenizer is None:
            model, tokenizer = load_fine_tuned_model()
        _LLM_WRAPPER = {"model": model, "tokenizer": tokenizer}
    return _LLM_WRAPPER
# Cached controls and document cache
document_cache = {}
document_cache_lock = threading.Lock()
cached_controls = None
BASE_DIR = os.environ.get('COMPLIANCE_BASE_DIR', '/home/gpu/Documents/Amal/Compliance_llama3.2')
MODEL_CACHE_DIR = os.environ.get('COMPLIANCE_MODEL_CACHE', os.path.join(BASE_DIR,'model_cache', 'finetuned_llama3.2'))
CHROMA_STORAGE_DIR = os.environ.get('COMPLIANCE_CHROMA_DIR', os.path.join(BASE_DIR, 'chroma_storage'))
CONTROLS_FILE = '/home/gpu/Documents/Amal/Compliance_llama3.2/control/controlt.json'

from difflib import SequenceMatcher

class ProgressWrapper:
    """A simple wrapper around Gradio's progress function to handle version differences."""
    
    def __init__(self, progress_obj=None):
        self.progress_obj = progress_obj
        self.current_progress = 0.0
        self.desc = "Starting..."
    
    def update(self, progress=None, desc=None):
        """Update progress."""
        if progress is not None:
            self.current_progress = progress
        if desc is not None:
            self.desc = desc
            
        # Call the wrapped progress object if it exists
        if self.progress_obj is not None:
            try:
                # Try the standard call with both progress and desc
                self.progress_obj(self.current_progress, self.desc)
            except Exception as e:
                try:
                    # Fall back to just progress value if that works
                    self.progress_obj(self.current_progress)
                    print(f"Progress message: {self.desc}")
                except Exception as e2:
                    # If all else fails, just print the message
                    print(f"Progress: {self.current_progress*100:.1f}% - {self.desc}")
        else:
            # No progress object, just print to console
            print(f"Progress: {self.current_progress*100:.1f}% - {self.desc}")
class ConsistentAnalysisCoordinator:
    """
    A coordinator class to maintain consistency across control analysis runs.
    """
    def __init__(self, seed=42):
        self.seed = seed
        self.control_seeds = {}  # Maps control IDs to their specific seeds
        self.cache = {}  # Simple cache of results
        self.base_seed = seed
    
    def get_control_seed(self, control_id):
        """Get a consistent seed for a specific control ID."""
        if control_id not in self.control_seeds:
            # Create a deterministic but different seed for each control
            # by hashing the control ID with the base seed
            control_seed = hash(f"{control_id}_{self.base_seed}") % 10000
            self.control_seeds[control_id] = control_seed
        return self.control_seeds[control_id]
    
    def analyze_control(self, control, document_text, embedding_model):
        """Perform consistent analysis of a control against a document."""
        control_id = control.get("id", "unknown")
        
        # Create a cache key based on control ID and document hash
        doc_hash = hash(document_text[:1000] + document_text[-1000:])
        cache_key = f"{control_id}_{doc_hash}"
        
        # Return cached result if available
        if cache_key in self.cache:
            return self.cache[cache_key]
        
        # Get a consistent seed for this control
        control_seed = self.get_control_seed(control_id)
        
        # Set global seed for this analysis
        torch.manual_seed(control_seed)
        if torch.cuda.is_available():
            torch.cuda.manual_seed_all(control_seed)
        np.random.seed(control_seed)
        
        # USE THE ORIGINAL WORKING FUNCTION
        try:
            result = analyze_control_agent(  # <-- CHANGE THIS BACK TO ORIGINAL
                control, 
                document_text, 
                embedding_model, 
                seed=control_seed
            )
            
            # Cache the result
            self.cache[cache_key] = result
            return result
            
        except Exception as e:
            print(f"Error in consistent analysis: {e}")
            # Return a consistent error result
            error_result = {
                "item": control,
                "type": control.get("type", "control"),
                "id": control_id,
                "description": control.get("description", ""),
                "status": "Not Addressed",
                "reason": f"Analysis error: {str(e)}",
                "evidence": "None",
                "full_response": ""
            }
            return error_result

class ContextualRetriever:
    """A pure embedding-based contextual retrieval system with caching."""
    
    # Class-level cache for document embeddings
    _document_embeddings_cache = {}
    
    def __init__(self, document_text, embedding_model):
        """Initialize with document text and a shared embedding model."""
        self.document_text = document_text
        self.embedding_model = embedding_model

        if self.embedding_model is None:
            raise ValueError("Embedding model is required for contextual retrieval.")

        # Cache key based on document content
        doc_hash = hash(document_text[:1000] + document_text[-1000:])
        
        # Check if this document is already in cache
        if doc_hash in self._document_embeddings_cache:
            write_debug_log(f"[DEBUG] Using cached sections and embeddings")
            self.sections, self.section_embeddings = self._document_embeddings_cache[doc_hash]
        else:
            # Extract sections using better section detection
            self.sections = self.extract_document_sections(document_text)
            
            # If section extraction fails, fall back to simple paragraph splitting
            if not self.sections:
                self.paragraphs = [p.strip() for p in document_text.split('\n\n') if p.strip()]
                self.sections = []
                section_size = 3

                for i in range(0, len(self.paragraphs), max(1, section_size - 1)):
                    section_paragraphs = self.paragraphs[i:i + section_size]
                    if section_paragraphs:
                        self.sections.append('\n\n'.join(section_paragraphs))

                if not self.sections and self.paragraphs:
                    self.sections = self.paragraphs

            # Compute section embeddings
            try:
                write_debug_log(f"[DEBUG] Computing embeddings for {len(self.sections)} sections")
                batch_size = 32
                all_embeddings = []

                for i in range(0, len(self.sections), batch_size):
                    batch = self.sections[i:i + batch_size]
                    batch_embeddings = self.embedding_model.encode(batch, convert_to_tensor=True)
                    all_embeddings.append(batch_embeddings.cpu().numpy())

                self.section_embeddings = np.vstack(all_embeddings)
                write_debug_log(f"[DEBUG] Embeddings computed successfully")
                
                # Cache for future use
                self._document_embeddings_cache[doc_hash] = (self.sections, self.section_embeddings)

            except Exception as e:
                write_debug_log(f"[ERROR] Failed during embedding: {e}")
                traceback.print_exc()
                self.section_embeddings = None

    def extract_document_sections(self, document_text):
        """Extract document sections using flexible formatting detection without hard-coded patterns."""
        try:
            # Start with the whole document as one section
            sections = [document_text]
            
            # Try multiple sectioning approaches and pick the one that gives the best results
            sectioning_approaches = []
            
            # Approach 1: Split by blank lines (paragraphs)
            paragraphs = [p.strip() for p in re.split(r'\n\s*\n', document_text) if p.strip()]
            if len(paragraphs) >= 5:
                sectioning_approaches.append(paragraphs)
            
            # Approach 2: Split by any line that looks like a header
            # This matches any line that's shorter than typical paragraphs and has special formatting
            header_pattern = r'\n(?:[^\n]{1,100}(?:\n\s*\n|\n\s*[A-Z0-9]))'
            header_splits = re.split(header_pattern, document_text)
            if len(header_splits) >= 3:
                sectioning_approaches.append(header_splits)
            
            # Approach 3: Look for any numerical or bullet-point structure
            # This catches numbered sections, bullet points, etc.
            structured_pattern = r'\n(?:(?:\d+\.|\*|\-|\•)\s+[^\n]+\n)'
            structured_splits = re.split(structured_pattern, document_text)
            if len(structured_splits) >= 3:
                sectioning_approaches.append(structured_splits)
                
            # Select the approach that gives the most balanced sections
            if sectioning_approaches:
                # Score each approach based on number of sections and their size distribution
                best_approach = None
                best_score = 0
                
                for approach in sectioning_approaches:
                    # Filter out very small sections
                    filtered_sections = [s for s in approach if len(s) > 50]
                    
                    if not filtered_sections:
                        continue
                        
                    # Calculate section length statistics
                    lengths = [len(s) for s in filtered_sections]
                    avg_length = sum(lengths) / len(lengths)
                    
                    # Calculate standard deviation (lower is better - more uniform sections)
                    variance = sum((l - avg_length) ** 2 for l in lengths) / len(lengths)
                    std_dev = variance ** 0.5
                    
                    # Ideal: More sections with lower standard deviation (more uniform)
                    # but don't penalize too much for fewer sections if they're consistent
                    score = len(filtered_sections) * (1000 / (std_dev + 1))
                    
                    if score > best_score:
                        best_score = score
                        best_approach = filtered_sections
                
                if best_approach:
                    sections = best_approach
            
            # Ensure we don't have too many tiny sections or too few giant sections
            # Consolidate small adjacent sections
            if len(sections) > 5:
                consolidated = []
                current = ""
                current_length = 0
                target_min_length = 300  # Aim for sections of at least 300 chars
                
                for section in sections:
                    if current_length < target_min_length:
                        current += "\n\n" + section
                        current_length += len(section)
                    else:
                        consolidated.append(current.strip())
                        current = section
                        current_length = len(section)
                
                if current:
                    consolidated.append(current.strip())
                
                # Only use consolidated if it gave us good results
                if consolidated and len(consolidated) >= 3:
                    sections = consolidated
            
            # Final check to avoid massive sections
            if len(sections) <= 2 and any(len(s) > 10000 for s in sections):
                # Just split very long sections by paragraph
                new_sections = []
                for section in sections:
                    if len(section) > 10000:
                        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', section) if p.strip()]
                        # Group paragraphs into manageable chunks
                        chunk = ""
                        for para in paragraphs:
                            if len(chunk) + len(para) < 5000:
                                chunk += "\n\n" + para
                            else:
                                new_sections.append(chunk.strip())
                                chunk = para
                        if chunk:
                            new_sections.append(chunk.strip())
                    else:
                        new_sections.append(section)
                if new_sections:
                    sections = new_sections
                    
            # Return non-empty sections only
            return [s for s in sections if s.strip()]
            
        except Exception as e:
            write_debug_log(f"[ERROR] Section extraction failed: {e}")
            return [document_text]
        
    def retrieve_context(self, query_text, control_id=None):
        """Context retrieval using purely semantic similarity without fixed patterns."""
        if self.section_embeddings is None:
            write_debug_log("[ERROR] Embedding model not available")
            return self.document_text[:min(6000, len(self.document_text))]

        try:
            # Build a semantically rich query
            context_terms = []
            if control_id:
                context_terms.append(control_id)
            context_terms.append(query_text)
            
            # Add compliance-related terminology to enhance query
            generic_terms = ["compliance", "policy", "requirement", "control", "standard"]
            enhanced_query = " ".join(context_terms + generic_terms)
            
            # Get query embedding
            query_embedding = self.embedding_model.encode([enhanced_query], convert_to_tensor=True)[0]
            query_embedding = query_embedding.cpu().numpy()
            
            # Calculate semantic similarity with all sections
            similarities = cosine_similarity([query_embedding], self.section_embeddings)[0]
            
            # Simple approach: Just get top sections by similarity score
            # This is purely semantic and not dependent on document structure
            sorted_indices = np.argsort(-similarities)
            top_sections = []
            total_chars = 0
            max_chars = 6000
            
            # Take sections until we hit the character limit
            for i in sorted_indices:
                if similarities[i] > 0.2:  # Only take somewhat relevant sections
                    if total_chars + len(self.sections[i]) <= max_chars:
                        top_sections.append((i, self.sections[i]))
                        total_chars += len(self.sections[i])
                        write_debug_log(f"[DEBUG] Added section {i} with similarity {similarities[i]:.4f}")
                        
                        # Stop if we have enough context with good quality
                        if similarities[i] < 0.25 and total_chars > 3000:
                            break
                            
                        # Stop if we have hit our overall limit
                        if total_chars >= 6000:
                            break
            
            # If we got very little content, lower the threshold and try again
            if total_chars < 1000:
                for i in sorted_indices:
                    if i not in [idx for idx, _ in top_sections]:
                        if total_chars + len(self.sections[i]) <= max_chars:
                            top_sections.append((i, self.sections[i]))
                            total_chars += len(self.sections[i])
                            
                            if total_chars >= 4000:
                                break
            
            # Format as context with section markers
            context = ""
            for idx, section in top_sections:
                # Add section with its similarity score (useful debugging info)
                context += f"\n\n--- SECTION (Similarity: {similarities[idx]:.2f}) ---\n{section}\n"
            
            write_debug_log(f"[DEBUG] Retrieved {len(top_sections)} sections ({len(context)} chars)")
            return context
            
        except Exception as e:
            write_debug_log(f"[ERROR] Retrieval error: {e}")
            traceback.print_exc()
            return self.document_text[:min(6000, len(self.document_text))]  
 
class ControlAnalysisResult(BaseModel):
    """Structured result for UAE IA compliance control analysis"""
    control_id: str = Field(description="Unique control identifier")
    status: Literal["Fully Addressed", "Partially Addressed", "Not Addressed"] = Field(
        description="Compliance status of the control"
    )
    reason: str = Field(description="Detailed explanation of the assessment")
    evidence: str = Field(description="Specific text evidence from documents as a simple string")
    confidence_score: Optional[float] = Field(
        default=None, 
        description="Confidence level between 0.0 and 1.0",
        ge=0.0, le=1.0
    )

    # ===============================
# 2. LLaMA Model Loader + Inference
# ===============================
   
def load_fine_tuned_model():
    global GLOBAL_MODEL, GLOBAL_TOKENIZER

    if GLOBAL_MODEL is not None and GLOBAL_TOKENIZER is not None:
        print("Using cached model and tokenizer")
        return GLOBAL_MODEL, GLOBAL_TOKENIZER

    try:
        model_path = "/home/gpu/Documents/Amal/Compliance_llama3.2/model_cache/finetuned_llama3.2"
        print(f"Loading fine-tuned Llama 3.2 model from {model_path}")
 
        tokenizer = AutoTokenizer.from_pretrained("meta-llama/Llama-3.2-1B-Instruct")
        tokenizer.pad_token = tokenizer.eos_token
        tokenizer.padding_side = 'left'
        
        # Load the base model
        base_model = AutoModelForCausalLM.from_pretrained(
            "meta-llama/Llama-3.2-1B-Instruct",
            torch_dtype=torch.float16,
            device_map="auto",
            
            trust_remote_code=True
        )
        
        # Try to load and merge the PEFT adapter
        try:
            model = PeftModel.from_pretrained(base_model, model_path)
            # Merge adapter with base model to avoid dimension issues
            model = model.merge_and_unload()
            print("PEFT adapter merged with base model")
        except Exception as e:
            print(f"Error loading PEFT adapter, using base model: {e}")
            model = base_model
        
        model.eval()
        
        # Set generation config
        model.config.pad_token_id = tokenizer.pad_token_id
        
        # Check if model loaded correctly
        if model is None:
            print("ERROR: Model failed to load properly")
            return None, None

        # Add model validation check
        try:
            with torch.no_grad():  # Disable gradients for efficiency
                # Just do a simple forward pass instead of generation
                input_ids = tokenizer("Test", return_tensors="pt").input_ids.to(model.device)
                _ = model(input_ids)
                print("Model loaded and validated successfully")
        except Exception as e:
            print(f"Model validation failed: {str(e)}")
            traceback.print_exc()
            return None, None
        
        # Only set global variables if all checks pass
        GLOBAL_MODEL = model
        GLOBAL_TOKENIZER = tokenizer
        print("Model loaded and ready.")
        return model, tokenizer
        
    except Exception as e:
        print(f"Error loading model: {e}")
        traceback.print_exc()
        return None, None

def get_llm_response(prompt, model, tokenizer, seed=42, max_retries=2):
    """Get consistent JSON responses from Llama 3.2 for Pydantic parsing."""
    torch.manual_seed(seed)
    if torch.cuda.is_available():
        torch.cuda.manual_seed_all(seed)
    
    write_debug_log(f"[DEBUG] Starting LLM generation with seed={seed}")
    
    # CREATE PIPELINE ONCE (better performance)
    pipe = pipeline("text-generation", model=model, tokenizer=tokenizer)
    
    for attempt in range(max_retries + 1):
        try:
            write_debug_log(f"[DEBUG] LLM attempt {attempt+1} with temperature=0.1")
            
            # Optimized config for JSON generation
            generation_config = {
                "max_new_tokens": 300,  # Shorter for JSON responses
                "do_sample": True,
                "temperature": 0.1,  # Lower for more consistent JSON
                "top_p": 0.9,  # Slightly lower
                "top_k": 20,   # More focused
                "repetition_penalty": 1.05,  # Less aggressive
                "pad_token_id": tokenizer.pad_token_id,
                "eos_token_id": tokenizer.eos_token_id,
                "return_full_text": False
            }
            
            # Add JSON instruction to prompt
            json_prompt = f"""{prompt}

IMPORTANT: Respond with ONLY a valid JSON object. No explanations, no markdown, no code blocks."""
            
            write_debug_log(f"[DEBUG] Calling LLM generation pipeline")
            start_time = time.time()
            
            result = pipe(json_prompt, **generation_config)
            response = result[0]["generated_text"].strip()
            
            generation_time = time.time() - start_time
            write_debug_log(f"[DEBUG] LLM generation completed in {generation_time:.2f} seconds")
            write_debug_log(f"[DEBUG] Response length: {len(response)} chars")
            
            # Validate JSON structure
            if '{' in response and '}' in response and len(response) > 20:
                write_debug_log(f"[DEBUG] LLM generation successful")
                return response
            else:
                raise ValueError(f"Invalid JSON structure in response: {response[:100]}")
            
        except Exception as e:
            write_debug_log(f"[ERROR] LLM generation attempt {attempt+1} failed: {str(e)}")
            
            if attempt < max_retries:
                write_debug_log(f"[DEBUG] Retrying attempt {attempt+2}")
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                continue
            else:
                write_debug_log(f"[ERROR] All {max_retries+1} LLM attempts failed")
                # Return valid JSON fallback
                fallback_json = f"""{{
    "control_id": "unknown",
    "status": "Not Addressed", 
    "reason": "LLM generation failed: {str(e)[:100]}",
    "evidence": "",
    "confidence_score": null
}}"""
                write_debug_log(f"[DEBUG] Returning JSON fallback response")
                return fallback_json
def extract_text_from_pdf_content(pdf_content):
    """Extract text from PDF binary content."""
    text = ""
    with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_docx_content(docx_content):
    """Extract text from DOCX binary content."""
    doc = Document(io.BytesIO(docx_content))
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_document(file_path):
    """Extract text from a document file with caching."""
    with document_cache_lock:
        if file_path in document_cache:
            return document_cache[file_path]
    
    ext = os.path.splitext(file_path.lower())[1]
    try:
        if ext == '.pdf':
            with open(file_path, 'rb') as file:
                pdf_content = file.read()
            text = extract_text_from_pdf_content(pdf_content)
        elif ext == '.docx':
            with open(file_path, 'rb') as file:
                docx_content = file.read()
            text = extract_text_from_docx_content(docx_content)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
        
        # Check if we got any text
        if not text or len(text.strip()) < 10:
            raise ValueError(f"No meaningful text extracted from {file_path}")
       
        # Call this after extraction
        text = preprocess_document_text(text)
        with document_cache_lock:
            document_cache[file_path] = text
        return text
    except Exception as e:
        return f"Error reading document: {str(e)}"
    
# Add after extracting text
def preprocess_document_text(text):
    """Clean and normalize document text."""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Normalize section headings
    text = re.sub(r'([A-Z][A-Za-z\s]+:)', r'\n\n\1', text)
    # Clean up common OCR artifacts
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', text)
    return text

# ===== Schema for LangGraph Reasoning Agent =====

from typing import Literal
class ControlAnalysisState(BaseModel):
    control_id: str
    control_description: str
    document_context: str
    analysis: Optional[str] = None
    requirements_identified: Optional[List[str]] = None
    matching_content: Optional[List[str]] = None
    status_determination: Optional[str] = None
    final_status: Optional[str] = None
    reason: Optional[str] = None
    evidence: Optional[str] = None

def analyze_compliance_with_agent(file_path, controls, model, tokenizer, progress=None, batch_size=4, max_controls=None):
    """Analyze compliance of a document against controls using agent-based approach."""
    try:
        # Extract text from document
        doc_text = extract_text_from_document(file_path)
        filename = os.path.basename(file_path)
        
        if progress:
            progress(f"Loading embedding model for semantic search")
        
        # Load embedding model
        try:
            embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            if torch.cuda.is_available():
                embedding_model = embedding_model.to('cuda')
        except Exception as e:
            print(f"Error loading embedding model: {e}")
            embedding_model = None
            
        # Create an analysis coordinator for consistency
        coordinator = ConsistentAnalysisCoordinator()
        
        if progress:
            progress(f"Analyzing {len(controls)} controls in {filename}")
            
        # Limit controls if max_controls specified
        if max_controls and len(controls) > max_controls:
            controls = controls[:max_controls]
            
        # Process controls in batches
        all_results = []
        total_controls = len(controls)
        
        # Process controls in batches for better performance
        for i in range(0, total_controls, batch_size):
            batch = controls[i:min(i + batch_size, total_controls)]
            
            if progress:
                progress(f"Processing controls {i+1}-{min(i+batch_size, total_controls)} of {total_controls}")
            
            # Process each control in the batch
            batch_results = []
            for control in batch:
                result = analyze_control_agent(control, doc_text, embedding_model)  # Call directly
                batch_results.append(result)
                
            all_results.extend(batch_results)
            
            # Update progress
            if progress:
                progress(f"Completed {min(i+batch_size, total_controls)}/{total_controls} controls")
                
        # Categorize results
        implemented = []
        not_addressed = []
        partially_addressed = []
        
        for result in all_results:
            status = result.get("status", "Not Addressed")
            
            if status == "Fully Addressed":
                implemented.append(result)
            elif status == "Partially Addressed":
                partially_addressed.append(result)
            else:
                not_addressed.append(result)
                
        return implemented, not_addressed, partially_addressed
        
    except Exception as e:
        print(f"Error in analyze_compliance_with_agent: {str(e)}")
        traceback.print_exc()
        return [], [], []

def analyze_control_agent(control, document_text, embedding_model, seed=42):
    """UAE IA compliance analysis with JSON, YAML, regex fallback parsing, and robust validation."""


    def sanitize(text: str) -> str:
        return text.strip().strip('"').strip(": ").replace('\n', ' ').strip()

    def validate_output(status, reason, evidence):
        """Downgrade status if evidence is insufficient."""
        if status == "Fully Addressed" and (len(evidence) < 50 or "no relevant content" in evidence.lower()):
            return "Partially Addressed", "Evidence insufficient or vague.", evidence or "No substantial evidence quoted."
        return status, reason, evidence

    control_id = control.get("id", "unknown")
    control_name = control.get("control_name", "Unknown")
    description = control.get("description", "No description provided.")

    try:
        torch.manual_seed(seed)
        if torch.cuda.is_available():
            torch.cuda.manual_seed_all(seed)

        retriever = ContextualRetriever(document_text, embedding_model)
        relevant_context = retriever.retrieve_context(f"{control_id} {description}", control_id)

        write_debug_log(f"[CONTROL] Analyzing: {control_id} - {description}")
        write_debug_log(f"[CONTEXT] Retrieved context length: {len(relevant_context)} chars")

        model_tokenizer = get_llm_wrapper()
        model, tokenizer = model_tokenizer["model"], model_tokenizer["tokenizer"]

        prompt = f"""You are a UAE IA compliance auditor. Analyze the following control against the provided document sections.

Control ID: {control_id}
Control Name: {control_name}
Requirement: {description}

Document Sections to Analyze:
{relevant_context}

TASK: Compare the control requirement with the document content above. 

If the document contains specific text that addresses this control requirement, mark as "Fully Addressed" and quote the relevant text.
If the document partially mentions related topics, mark as "Partially Addressed" and explain what's missing.
If the document has no relevant content, mark as "Not Addressed".

Return ONLY this format:
control_id: "{control_id}"
status: "Fully Addressed" OR "Partially Addressed" OR "Not Addressed"
reason: Brief explanation of your analysis
evidence: Exact quotes from document OR "No relevant content found"

DO NOT repeat the control description. Analyze the actual document content."""

        response = get_llm_response(prompt, model, tokenizer, seed)
        write_debug_log(f"[RAW_RESPONSE] {control_id}: {response[:200]}...")

        cleaned_response = response.strip()
        for marker in ['Here is the response:', 'Here is the output:', 'Just the control ID']:
            if marker in cleaned_response:
                cleaned_response = cleaned_response.split(marker)[-1].strip()
        if cleaned_response.startswith('```'):
            cleaned_response = re.sub(r'^```[^\n]*\n?', '', cleaned_response)
            cleaned_response = re.sub(r'\n?```$', '', cleaned_response)

        # Try JSON with control ID as key
        try:
            parsed_json = json.loads(cleaned_response)
            if control_id in parsed_json:
                data = parsed_json[control_id]
                status, reason, evidence = validate_output(
                    sanitize(data.get("status", "Not Addressed")),
                    sanitize(data.get("reason", "")),
                    sanitize(data.get("evidence", ""))
                )
                return {
                    "item": control,
                    "type": control.get("type", "control"),
                    "id": control_id,
                    "description": description,
                    "status": status,
                    "reason": reason,
                    "evidence": evidence,
                    "full_response": response
                }
        except json.JSONDecodeError as e:
            write_debug_log(f"[JSON_ERROR] {str(e)}")

        # Try YAML with control ID as key
        try:
            parsed_yaml = yaml.safe_load(cleaned_response)
            if isinstance(parsed_yaml, dict) and control_id in parsed_yaml:
                data = parsed_yaml[control_id]
                status, reason, evidence = validate_output(
                    sanitize(data.get("status", "Not Addressed")),
                    sanitize(data.get("reason", "")),
                    sanitize(data.get("evidence", ""))
                )
                return {
                    "item": control,
                    "type": control.get("type", "control"),
                    "id": control_id,
                    "description": description,
                    "status": status,
                    "reason": reason,
                    "evidence": evidence,
                    "full_response": response
                }
        except yaml.YAMLError as e:
            write_debug_log(f"[YAML_ERROR] {str(e)}")

        # Try JSON with "control_id" field
        try:
            parsed_alt = json.loads(cleaned_response)
            if isinstance(parsed_alt, dict) and parsed_alt.get("control_id") == control_id:
                status, reason, evidence = validate_output(
                    sanitize(parsed_alt.get("status", "Not Addressed")),
                    sanitize(parsed_alt.get("reason", "")),
                    sanitize(parsed_alt.get("evidence", ""))
                )
                return {
                    "item": control,
                    "type": control.get("type", "control"),
                    "id": control_id,
                    "description": description,
                    "status": status,
                    "reason": reason,
                    "evidence": evidence,
                    "full_response": response
                }
        except Exception as e:
            write_debug_log(f"[ALT_JSON_PARSE_ERROR] {str(e)}")

        # Regex fallback
        pattern = re.search(
            rf'{re.escape(control_id)}.*?'
            r'status[:\-]?\s*(?P<status>Fully Addressed|Partially Addressed|Not Addressed).*?'
            r'reason[:\-]?\s*(?P<reason>.+?)\s*'
            r'evidence[:\-]?\s*(?P<evidence>.+)',
            cleaned_response,
            re.IGNORECASE | re.DOTALL
        )

        if pattern:
            status, reason, evidence = validate_output(
                sanitize(pattern.group("status")),
                sanitize(pattern.group("reason")),
                sanitize(pattern.group("evidence"))
            )
            write_debug_log(f"[REGEX_PARSE] Fallback regex extraction succeeded for {control_id}")
            return {
                "item": control,
                "type": control.get("type", "control"),
                "id": control_id,
                "description": description,
                "status": status,
                "reason": reason,
                "evidence": evidence,
                "full_response": response
            }

        write_debug_log(f"[MANUAL_ERROR] Could not extract control info for {control_id}")
        return {
            "item": control,
            "type": control.get("type", "control"),
            "id": control_id,
            "description": description,
            "status": "Not Addressed",
            "reason": "No analysis available...",
            "evidence": "No evidence found...",
            "full_response": response
        }

    except Exception as e:
        write_debug_log(f"[ERROR] Critical error for {control_id}: {e}")
        return {
            "item": control,
            "type": control.get("type", "control"),
            "id": control_id,
            "description": description,
            "status": "Not Addressed",
            "reason": f"Critical error: {str(e)}",
            "evidence": "No evidence available",
            "full_response": ""
        }


def load_controls():
    """Load controls from JSON files with improved error handling"""
    global cached_controls
    
    if cached_controls is not None:
        return cached_controls
    all_controls=[]    
    try:
        control_json_path = CONTROLS_FILE
        print(f"Loading controls from: {control_json_path}")
        
        if not os.path.exists(control_json_path):
            print(f"Control file not found: {control_json_path}")
            return []
        
        with open(control_json_path, "r", encoding='utf-8') as f:
            content = f.read()
            
            # Fix JSON formatting issues
            content = re.sub(r',\s*]', ']', content)
            content = re.sub(r',\s*}', '}', content)
            
            try:
                control_data = json.loads(content)
            except json.JSONDecodeError as e:
                print(f"JSON parsing failed: {e}")
                control_data = parse_json_objects(content)
                
            if isinstance(control_data, dict) and "controls" in control_data:
                controls = control_data["controls"]
            elif isinstance(control_data, list):
                controls = control_data
            else:
                print(f"Unexpected control data structure: {type(control_data)}")
                return []
            
            for control in controls:
                mapped_control = {
                    "type": control.get("control_sub_family", "Unknown"),
                    "id": control.get("control_code", "Unknown"),
                    "content": control.get("control_name", ""),
                    "section": control.get("section", ""),
                    "control_family": control.get("control_family", ""),
                    "control_sub_family": control.get("control_sub_family", ""),
                    "control_code": control.get("control_code", ""),
                    "control_name": control.get("control_name", ""),
                    "priority": control.get("priority", ""),
                    "applicability": control.get("applicability", ""),
                    "description": control.get("description", "")
                }
                all_controls.append(mapped_control)
    
        control_types = {}
        for item in all_controls:
            item_type = item.get("type", "control")
            if item_type not in control_types:
                control_types[item_type] = 0
            control_types[item_type] += 1
    
        print(f"Controls by type: {control_types}")
        print(f"Total controls: {len(all_controls)}")
        
        cached_controls = all_controls
        return all_controls
        
    except Exception as e:
        print(f"Error loading controls: {str(e)}")
        traceback.print_exc()
        return []

def parse_json_objects(text):
    """Parse JSON objects from text, even if malformed"""
    controls = []
    pattern = r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}'
    matches = re.findall(pattern, text)
    
    for match in matches:
        try:
            obj = json.loads(match)
            controls.append(obj)
        except json.JSONDecodeError:
            continue
    
    return controls

def process_upload(file_path, progress=None):
    """Process a single uploaded file and run compliance checks."""
    filename = os.path.basename(file_path)
    
    if progress:
        progress(f"Extracting text from {filename}")
        
    # Extract text
    doc_text = extract_text_from_document(file_path)
    
    if isinstance(doc_text, str) and doc_text.startswith("Error"):
        summary = f"<div style='background-color: #f8d7da; color: #721c24; padding: 15px; border-radius: 5px;'><h3>Error Processing File</h3><p>{doc_text}</p></div>"
        return summary, "", "", "", [], [], [], filename
    
    if progress:
        progress(f"Loading model and controls")
        
    # Load model and tokenizer
    model, tokenizer = load_fine_tuned_model()
    if model is None or tokenizer is None:
        summary = "<div style='background-color: #f8d7da; color: #721c24; padding: 15px; border-radius: 5px;'><h3>Error</h3><p>Failed to load the fine-tuned model and tokenizer.</p></div>"
        return summary, "", "", "", [], [], [], filename
    
      
    # Load controls
    print(f"[DEBUG] Loading controls")
    write_debug_log(f"[DEBUG] Loading controls")
    controls = load_controls()
    if not controls:
        summary = "<div style='background-color: #f8d7da; color: #721c24; padding: 15px; border-radius: 5px;'><h3>Error</h3><p>Failed to load control files.</p></div>"
        return summary, "", "", "", [], [], [], filename
    
    # Analyze compliance using parallel processing
    implemented, not_addressed, partially_addressed = analyze_compliance_with_agent(
        file_path, 
        controls, 
        model, 
        tokenizer,
        progress=lambda msg: progress(msg) if progress else None,
        batch_size=4,
        max_controls=25
    )
    not_applicable = []
    
    if progress:
        progress(f"Generating results")
    
    # Create summary output
    total_items = len(implemented) + len(not_addressed) + len(partially_addressed)
    
    if total_items == 0:
        summary = "<div style='background-color: #f8d7da; color: #721c24; padding: 15px; border-radius: 5px;'><h3>No Results</h3><p>No compliance items could be analyzed. There might be an issue with the document format or content.</p></div>"
        return summary, "", "", "", [], [], [], filename
    
    summary = f"""
    <div style='background-color: #f0f5ff; padding: 20px; border-radius: 5px; border-left: 4px solid #3498db; margin-bottom: 20px;'>
        <h2>Compliance Analysis Results for: {filename}</h2>
        <div style='margin-top: 20px; padding: 15px; background-color: #e8f4f8; border-radius: 5px;'>
            <h3>Compliance Status</h3>
            <div style='display: flex; flex-wrap: wrap;'>
                <div style='margin-right: 30px;'>
                    <div style='color: #7f8c8d; font-size: 14px;'>Total items checked</div>
                    <div style='font-size: 24px; font-weight: bold;'>{total_items}</div>
                </div>
                <div style='margin-right: 30px;'>
                    <div style='color: #7f8c8d; font-size: 14px;'>Fully Addressed</div>
                    <div style='font-size: 24px; font-weight: bold; color: #27ae60;'>{len(implemented)} ({len(implemented)/total_items*100:.1f}%)</div>
                </div>
                <div style='margin-right: 30px;'>
                    <div style='color: #7f8c8d; font-size: 14px;'>Partially Addressed</div>
                    <div style='font-size: 24px; font-weight: bold; color: #f39c12;'>{len(partially_addressed)} ({len(partially_addressed)/total_items*100:.1f}%)</div>
                </div>
                <div style='margin-right: 30px;'>
                    <div style='color: #7f8c8d; font-size: 14px;'>Controls Not Addressed</div>
                    <div style='font-size: 24px; font-weight: bold; color: #e74c3c;'>{len(not_addressed)} ({len(not_addressed)/total_items*100:.1f}%)</div>
                </div>
            </div>
        </div>
    </div>
    """
    
    # Format the outputs
    implemented_text = format_output_for_gradio_html(implemented, "Fully Addressed")
    missing_text = format_output_for_gradio_html(not_addressed, "Controls Not Addressed")
    partially_text = format_output_for_gradio_html(partially_addressed, "Partially Addressed")
    
    return summary, implemented_text, missing_text, partially_text, implemented, not_addressed, partially_addressed, filename
    

def format_output_for_gradio_html(items, status_type):
    """Format items for display in Gradio as HTML tables - optimized for performance."""
    if not items:
        return f"<div style='background-color: #f8d7da; color: #721c24; padding: 15px; border-radius: 5px;'>No {status_type} found.</div>"
    # Clean up the items before displaying
    cleaned_items = []
    for item in items:
        clean_item = item.copy()  # Make a copy to avoid modifying the original
        
        # Clean description - use the one from control JSON
        clean_item["description"] = clean_item.get("description", "")
        if len(clean_item["description"]) > 150:
            clean_item["description"] = clean_item["description"][:150] + "..."
        
        # Clean reason
        reason = clean_item.get("reason", "")
        # Remove underscores and excessive whitespace
        reason = re.sub(r'_{3,}', '', reason)
        reason = re.sub(r'\s+', ' ', reason).strip()
        # Remove prompt elements
        reason = re.sub(r'Step \d+:|TASK:|Answer:|The final answer is:', '', reason)
        reason = re.sub(r'\$\\boxed\{.*?\}\$', '', reason)
        if len(reason) > 150:
            reason = reason[:150] + "..."
        clean_item["reason"] = reason
        
        # Clean evidence
        evidence = clean_item.get("evidence", "")
        # Remove placeholder text and format elements
        evidence = re.sub(r'TASK: If the status.*', '', evidence)
        evidence = re.sub(r'_{3,}', '', evidence)
        evidence = re.sub(r'\s+', ' ', evidence).strip()
        if "None" in evidence and len(evidence) < 10:
            evidence = "No specific evidence found"
        elif len(evidence) > 150:
            evidence = evidence[:150] + "..."
        clean_item["evidence"] = evidence
        
        cleaned_items.append(clean_item)
    items_by_type = {}
    for item in items:
        item_type = item.get("type", "Unknown")
        if item_type not in items_by_type:
            items_by_type[item_type] = []
        items_by_type[item_type].append(item)
    
    output = "<style>\n"
    output += "table { width: 100%; border-collapse: collapse; margin-bottom: 20px; }\n"
    output += "th { background-color: #2c3e50; color: white !important; text-align: left; padding: 12px; }\n"
    output += "td { padding: 10px; border-bottom: 1px solid #ddd; }\n"
    output += "tr:nth-child(even) { background-color: #f2f2f2; }\n"
    output += "tr:hover { background-color: #e3f2fd; }\n"
    output += ".item-header { background-color: #3498db; color: white !important; padding: 10px; margin-top: 20px; border-radius: 5px; }\n"
    output += ".status-addressed { color: #27ae60; font-weight: bold; }\n"
    output += ".status-not_addressed { color: #e74c3c; font-weight: bold; }\n"
    output += ".status-partially { color: #f39c12; font-weight: bold; }\n"
    output += ".evidence { background-color: #ecf0f1; padding: 10px; border-left: 4px solid #3498db; margin: 10px 0; }\n"
    output += ".reason { margin-bottom: 10px; }\n"
    output += "</style>\n"
    
    MAX_ITEMS_PER_TYPE = 30  
    for type_name, type_items in items_by_type.items():
        output += f"<div class='item-header'><h2>{type_name.capitalize()}s ({len(type_items)})</h2></div>\n"
        
        output += "<table>\n"
        output += "<tr><th>ID</th><th>Description</th><th>Status</th></tr>\n"
        
        items_to_show = type_items[:MAX_ITEMS_PER_TYPE]
        for i, item in enumerate(items_to_show):
            output += "<tr>"
            output += f"<td>{item['id']}</td>"
            
            desc = item['description']
            if len(desc) > 150:
                desc = desc[:150] + "..."
            
            details = f"<div>{desc}</div>"
            if 'reason' in item and item['reason']:
                reason = item['reason']
                # Clean up technical error messages
                if "Error during analysis:" in reason:
                    reason = "Unable to analyze control"
                elif "shapes cannot be multiplied" in reason:
                    reason = "Technical error in processing"
                
                if len(reason) > 150:
                    reason = reason[:150] + "..."
                details += f"<div class='reason'><strong>Reason:</strong> {reason}</div>"
                
            if 'evidence' in item and item['evidence']:
                evidence = item['evidence']
                if len(evidence) > 150:
                    evidence = evidence[:150] + "..."
                details += f"<div class='evidence'><strong>Evidence:</strong> {evidence}</div>"
                
            output += f"<td>{details}</td>"
            
            status = item.get('status', 'Unknown')
            status_class = ""
            if status == "Implemented" or status == "Fully Addressed":
                status_class = "status-addressed"
            elif status == "Partially Addressed":
                status_class = "status-partially"
            elif status == "Not Implemented" or status == "Not Addressed":
                status_class = "status-not_addressed"
            output += f"<td><span class='{status_class}'>{status}</span></td>"
            
            output += "</tr>\n"
        
        if len(type_items) > MAX_ITEMS_PER_TYPE:
            remaining = len(type_items) - MAX_ITEMS_PER_TYPE
            output += f"<tr><td colspan='3'>Showing {MAX_ITEMS_PER_TYPE} of {len(type_items)} items. {remaining} more items are available in the downloaded report.</td></tr>"
            
        output += "</table>\n"
    
    return output

def process_file_worker(file, model, tokenizer, controls):
    """Worker function to process a single file."""
    try:
        filename = os.path.basename(file)
        print(f"Processing file: {filename}")
        
        # Extract text (with caching)
        doc_text = extract_text_from_document(file)
        
        # Check for extraction errors
        if isinstance(doc_text, str) and doc_text.startswith("Error"):
            print(f"Error extracting text from {filename}: {doc_text}")
            return {
                "filename": filename,
                "implemented": [],
                "missing": [],
                "partially_addressed": [],
                "error": doc_text
            }
        
        # Analyze with agent-based approach but without progress parameter
        implemented, not_addressed, partially_addressed = analyze_compliance_with_agent(
            file, controls, model, tokenizer, batch_size=4, max_controls=25
        )
        
        # Fix status mapping before returning
        for item in implemented:
            item["status"] = "Fully Addressed"
        
        for item in not_addressed:
            item["status"] = "Not Addressed"
            
        for item in partially_addressed:
            item["status"] = "Partially Addressed"
        
        return {
            "filename": filename,
            "implemented": implemented,
            "missing": not_addressed,
            "partially_addressed": partially_addressed,
            "error": None
        }
    except Exception as e:
        print(f"Error processing file {os.path.basename(file)}: {str(e)}")
        traceback.print_exc()
        return {
            "filename": os.path.basename(file),
            "implemented": [],
            "missing": [],
            "partially_addressed": [],
            "error": str(e)
        }
def process_multiple_uploads(files):
    """Process multiple uploaded files without progress bar."""
    if not files or len(files) == 0:
        return (
            "<p>Please upload at least one document file.</p>", 
            "<p>No files selected.</p>", 
            "<p>No files selected.</p>",
            "<p>No files selected.</p>",
            [], [], [], [],
            "",
            []
        )
    
    print("Loading models and controls...")
    
    model, tokenizer = load_fine_tuned_model()
    if model is None or tokenizer is None:
        return (
            "<p>Error loading fine-tuned model. Please check the model path.</p>",
            "<p>Model loading error</p>",
            "<p>Model loading error</p>",
            "<p>Model loading error</p>",
            [], [], [], [],
            "",
            []
        )
    
    all_controls = load_controls()
    if not all_controls:
        return (
            "<p>Error loading control files. Please check if the control files exist.</p>", 
            "<p>Error</p>", 
            "<p>Error</p>",
            "<p>Error</p>",
            [], [], [], [],
            "",
            []
        )
    
    print(f"Starting analysis for {len(files)} files...")
    
    # Process files in sequential manner to avoid parallel processing issues
    all_results = []
    file_count = len(files)
    
    for i, file_path in enumerate(files):
        print(f"Processing file {i+1}/{file_count}: {os.path.basename(file_path)}")
        
        try:
            # Process file (without progress parameter)
            result = process_file_worker(file_path, model, tokenizer, all_controls)
            all_results.append(result)
            
            total_controls = len(result.get('implemented', [])) + len(result.get('partially_addressed', [])) + len(result.get('missing', []))
            print(f"Completed file {i+1}/{file_count}: {result['filename']} ({total_controls} controls)")
            
        except Exception as e:
            print(f"Error processing file: {e}")
            traceback.print_exc()
    
    # Clear GPU cache
    if torch.cuda.is_available():
        torch.cuda.empty_cache()
    
    # Combine all results
    print("Combining results from all files...")
    
    total_implemented = []
    total_missing = []
    total_partially_addressed = []
    
    for result in all_results:
        filename = result.get("filename", "Unknown")
        for control in result.get("implemented", []):
            control["source_document"] = filename
        for control in result.get("partially_addressed", []):
            control["source_document"] = filename
        for control in result.get("missing", []):
            control["source_document"] = filename
        
        total_implemented.extend(result.get("implemented", []))
        total_missing.extend(result.get("missing", []))
        total_partially_addressed.extend(result.get("partially_addressed", []))
    
    print("Generating summary and reports...")
    
    # Create summary output
    summary = "<div style='background-color: #f0f5ff; padding: 20px; border-radius: 5px; border-left: 4px solid #3498db; margin-bottom: 20px;'>"
    summary += "<h2>Multi-Document Compliance Analysis Results</h2>"
    summary += "<div style='margin-bottom: 20px;'>Analyzed files: " + str(len(files)) + "</div>"
    summary += "<table style='width: 100%; border-collapse: collapse;'>"
    summary += "<tr style='background-color: #2c3e50; color: white;'><th style='padding: 10px; text-align: left;'>Document</th><th style='padding: 10px; text-align: center;'>Total</th><th style='padding: 10px; text-align: center;'>Controls Fully Addressed</th><th style='padding: 10px; text-align: center;'>Partially Addressed</th><th style='padding: 10px; text-align: center;'>Controls Not Addressed</th></tr>"
    
    # Add row for each file
    for result in all_results:
        filename = result.get("filename", "Unknown")
        implemented = len(result.get("implemented", []))
        partially_addressed = len(result.get("partially_addressed", []))
        not_addressed = len(result.get("missing", []))
        total_items = implemented + partially_addressed + not_addressed
        
        summary += f"<tr style='border-bottom: 1px solid #ddd;'>"
        summary += f"<td style='padding: 10px;'>{filename}</td>"
        summary += f"<td style='padding: 10px; text-align: center;'>{total_items}</td>"
        summary += f"<td style='padding: 10px; text-align: center; color: #27ae60;'>{implemented}</td>"
        summary += f"<td style='padding: 10px; text-align: center; color: #f39c12;'>{partially_addressed}</td>"
        summary += f"<td style='padding: 10px; text-align: center; color: #e74c3c;'>{not_addressed}</td>"
        summary += "</tr>"

    # Close summary table
    summary += "</table>"

    # Add overall statistics
    total_all = len(total_implemented) + len(total_missing) + len(total_partially_addressed)
    if total_all > 0:
        summary += "<div style='margin-top: 20px; padding: 15px; background-color: #e8f4f8; border-radius: 5px;'>"
        summary += "<h3>Overall Compliance Status</h3>"
        summary += "<div style='display: flex; flex-wrap: wrap;'>"
        summary += f"<div style='margin-right: 30px;'><div style='color: #7f8c8d; font-size: 14px;'>Total items checked</div><div style='font-size: 24px; font-weight: bold;'>{total_all}</div></div>"
        summary += f"<div style='margin-right: 30px;'><div style='color: #7f8c8d; font-size: 14px;'>Controls Fully Addressed</div><div style='font-size: 24px; font-weight: bold; color: #27ae60;'>{len(total_implemented)} ({len(total_implemented)/total_all*100:.1f}%)</div></div>"
        summary += f"<div style='margin-right: 30px;'><div style='color: #7f8c8d; font-size: 14px;'>Partially Addressed</div><div style='font-size: 24px; font-weight: bold; color: #f39c12;'>{len(total_partially_addressed)} ({len(total_partially_addressed)/total_all*100:.1f}%)</div></div>"
        summary += f"<div style='margin-right: 30px;'><div style='color: #7f8c8d; font-size: 14px;'>Controls Not Addressed</div><div style='font-size: 24px; font-weight: bold; color: #e74c3c;'>{len(total_missing)} ({len(total_missing)/total_all*100:.1f}%)</div></div>"
        summary += "</div></div>"

    summary += "</div>"

    # Format output for Gradio
    implemented_text = format_output_for_gradio_html(total_implemented, "Fully Addressed")
    partially_text = format_output_for_gradio_html(total_partially_addressed, "Partially Addressed")
    missing_text = format_output_for_gradio_html(total_missing, "Controls Not Addressed")

    total_count = len(total_implemented) + len(total_partially_addressed) + len(total_missing)
    print(f"✅ Analysis complete: {total_count} controls across {len(files)} files")

    return summary, implemented_text, partially_text, missing_text, total_implemented, total_partially_addressed, total_missing, "multiple_documents", all_results


def generate_combined_report_html(all_results):
    """Generate a comprehensive combined report from all results."""
    print(f"Combined report generating with {len(all_results)} document results")
    if not all_results or len(all_results) == 0:
        return "<p>No data available for combined report</p>"
    
    try:
        documents = [result.get('filename', 'Unknown') for result in all_results]
        
        html = """
        <style>
        .combined-report {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background-color: white;
            border-radius: 8px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
            padding: 25px;
            margin-top: 20px;
        }
        .report-header {
            background: linear-gradient(135deg, #155799, #159957);
            color: white !important;
            padding: 15px 20px;
            border-radius: 8px 8px 0 0;
            margin: -25px -25px 20px -25px;
        }
        .report-section {
            margin-bottom: 30px;
        }
        .report-section h3 {
            border-bottom: 2px solid #eaecef;
            padding-bottom: 8px;
            margin-top: 30px;
        }
        .document-section {
            background-color: #f8f9fa;
            border-radius: 6px;
            padding: 15px;
            margin-bottom: 20px;
            border-left: 4px solid #3498db;
        }
        .control-table {
            width: 100%;
            border-collapse: collapse;
            margin-top: 15px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.07);
        }
        .control-table th {
            background-color: #f8f9fa;
            color: #333;
            padding: 12px 15px;
            text-align: left;
            border: 1px solid #e9ecef;
        }
        .control-table td {
            padding: 12px 15px;
            border: 1px solid #e9ecef;
            vertical-align: top;
        }
        .status-implemented {
            color: #27ae60;
            font-weight: bold;
        }
        .status-partially {
            color: #f39c12;
            font-weight: bold;
        }
        .status-not-implemented {
            color: #e74c3c;
            font-weight: bold;
        }
        
        .evidence-block {
            background-color: #f8f9fa;
            border-left: 3px solid #3498db;
            padding: 10px;
            margin-top: 10px;
            font-size: 0.9em;
        }
        </style>
        
        <div class="combined-report">
            <div class="report-header">
                <h2>Comprehensive Compliance Analysis Report</h2>
                <p>Detailed analysis of all controls across all documents</p>
            </div>
        """
        # Now add the documents summary section
        html += """
        <div class="report-section">
            <h3>Documents Analyzed</h3>
            <table class="control-table" style="width:100%; border-collapse:collapse; margin-top:15px;">
                <tr>
                    <th>Filename</th>
                    <th>Controls Analyzed</th>
                </tr>
        """

        # Clean up data in all_results
        for result in all_results:
            for collection in ["implemented", "partially_addressed", "missing"]:
                for item in result.get(collection, []):
                    # Clean reason
                    reason = item.get("reason", "")
                    # Remove underscores and excessive whitespace
                    reason = re.sub(r'_{3,}', '', reason)
                    reason = re.sub(r'\s+', ' ', reason).strip()
                    # Remove prompt elements
                    reason = re.sub(r'Step \d+:|TASK:|Answer:|The final answer is:', '', reason)
                    reason = re.sub(r'\$\\boxed\{.*?\}\$', '', reason)
                    if len(reason) > 200:
                        reason = reason[:200] + "..."
                    item["reason"] = reason
                    
                    # Clean evidence
                    evidence = item.get("evidence", "")
                    # Remove placeholder text and format elements
                    evidence = re.sub(r'TASK: If the status.*', '', evidence)
                    evidence = re.sub(r'_{3,}', '', evidence)
                    evidence = re.sub(r'\s+', ' ', evidence).strip()
                    if "None" in evidence and len(evidence) < 10:
                        evidence = "No specific evidence found"
                    elif len(evidence) > 200:
                        evidence = evidence[:200] + "..."
                    item["evidence"] = evidence
        
        # Add rows for each document with total controls count
        for result in all_results:
            filename = result.get('filename', 'Unknown')
            total_controls = (
                len(result.get('implemented', [])) + 
                len(result.get('partially_addressed', [])) + 
                len(result.get('missing', []))
            )
            html += f"<tr><td>{filename}</td><td>{total_controls}</td></tr>"

        html += """
            </table>
        </div>
        """
        
        # Per-document detailed analysis
       
        for i, result in enumerate(all_results):
            filename = result.get('filename', 'Unknown')
            implemented = result.get('implemented', [])
            partially_addressed = result.get('partially_addressed', [])
            not_addressed = result.get('missing', [])
            
            html += f"""
            <div class="report-section">
                <h3>Detailed Analysis: {filename}</h3>
            """
            
            # Fully Addressed controls
            if implemented:
                html += """
                <h4 style="color: #27ae60;">Fully Addressed Controls</h4>
                <table class="control-table">
                    <tr><th>Control ID</th><th>Description</th><th>Reason</th><th>Evidence</th></tr>
                """
                
                for control in implemented:
                    html += f"""
                    <tr>
                        <td>{control.get('id', 'Unknown')}</td>
                        <td>{control.get('description', '')}</td>
                        <td>{control.get('reason', '')}</td>
                        <td>
                    """
                    
                    if control.get('evidence'):
                        html += f"""<div class="evidence-block">{control.get('evidence', '')}</div>"""
                    
                    html += """
                        </td>
                    </tr>
                    """
                
                html += "</table>"
            
            # Partially Addressed controls
            if partially_addressed:
                html += """
                <h4 style="color: #f39c12;">Partially Addressed Controls</h4>
                <table class="control-table">
                    <tr><th>Control ID</th><th>Description</th><th>Reason</th><th>Evidence</th></tr>
                """
                
                for control in partially_addressed:
                    html += f"""
                    <tr>
                        <td>{control.get('id', 'Unknown')}</td>
                        <td>{control.get('description', '')}</td>
                        <td>{control.get('reason', '')}</td>
                        <td>
                    """
                    
                    if control.get('evidence'):
                        html += f"""<div class="evidence-block">{control.get('evidence', '')}</div>"""
                    
                    html += """
                        </td>
                    </tr>
                    """
                
                html += "</table>"
            
            # Not addressed controls
            if not_addressed:
                html += """
                <h4 style="color: #e74c3c;">Controls Not Addressed</h4>
                <table class="control-table">
                    <tr><th>Control ID</th><th>Description</th><th>Reason</th></tr>
                """
                
                for control in not_addressed:
                    html += f"""
                    <tr>
                        <td>{control.get('id', 'Unknown')}</td>
                        <td>{control.get('description', '')}</td>
                        <td>{control.get('reason', '')}</td>
                    </tr>
                    """
                
                html += "</table>"
            html += "</div>"  # End document section     
                
        return html        
                
    except Exception as e:
        print(f"Error creating combined report: {str(e)}")
        traceback.print_exc()
        return f"<p>Error generating combined report: {str(e)}</p>"

def generate_matrix_report_html(all_results, all_controls=None):
    """Generate a matrix-style HTML report showing controls vs documents."""
    if not all_results or len(all_results) == 0:
        return "<p>No data available for the matrix report</p>"
    
    try:
        if all_controls is None:
            all_controls = {}
            for result in all_results:
                # Include all three categories of controls
                for implemented in result.get('implemented', []):
                    control_id = implemented.get('id')
                    control_type = implemented.get('type', 'Unknown')
                    if control_id not in all_controls:
                        all_controls[control_id] = {
                            'id': control_id,
                            'type': control_type,
                            'description': implemented.get('description', '')
                        }
                
                for partially_addressed in result.get('partially_addressed', []):
                    control_id = partially_addressed.get('id')
                    control_type = partially_addressed.get('type', 'Unknown')
                    if control_id not in all_controls:
                        all_controls[control_id] = {
                            'id': control_id,
                            'type': control_type,
                            'description': partially_addressed.get('description', '')
                        }
                
                for not_addressed in result.get('missing', []):
                    control_id = not_addressed.get('id')
                    control_type = not_addressed.get('type', 'Unknown')
                    if control_id not in all_controls:
                        all_controls[control_id] = {
                            'id': control_id,
                            'type': control_type,
                            'description': not_addressed.get('description', '')
                        }
        
        # Create the status matrix
        matrix = {}
        for control_id, control in all_controls.items():
            matrix[control_id] = {
                'control': control,
                'status': {}
            }
            for result in all_results:
                document = result.get('filename', 'Unknown')
                # Default to Not Addressed (Red)
                matrix[control_id]['status'][document] = {
                    'status': 'Controls Not Addressed',
                    'color': '#e74c3c',  # Red
                    'reason': '',
                    'evidence': ''
                }
                
                # Check if implemented (overwrite default if found)
                for implemented in result.get('implemented', []):
                    if implemented.get('id') == control_id:
                        matrix[control_id]['status'][document] = {
                            'status': 'Fully Addressed',
                            'color': '#27ae60',  # Green
                            'reason': implemented.get('reason', ''),
                            'evidence': implemented.get('evidence', '')
                        }
                        break
                
                # Check if partially addressed (overwrite default if found)
                for partial in result.get('partially_addressed', []):
                    if partial.get('id') == control_id:
                        matrix[control_id]['status'][document] = {
                            'status': 'Partially Addressed',
                            'color': '#f39c12',  # Orange
                            'reason': partial.get('reason', ''),
                            'evidence': partial.get('evidence', '')
                        }
                        break

        # Group controls by type
        controls_by_type = {}
        for control_id, data in matrix.items():
            control_type = data['control'].get('type', 'Unknown')
            if control_type not in controls_by_type:
                controls_by_type[control_type] = []
            controls_by_type[control_type].append(data)
        
        # Get document names
        document_names = []
        print(f"Matrix report generating with {len(all_results)} document results")
        for result in all_results:
            document_names.append(result.get('filename', 'Unknown'))
            print(f"Document: {result.get('filename')}, Controls: {len(result.get('implemented', [])) + len(result.get('partially_addressed', [])) + len(result.get('missing', []))}")
        
        # Create the HTML
        html = """
        <style>
        .matrix-container {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background-color: white;
            border-radius: 8px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
            padding: 25px;
            margin-top: 20px;
        }
        .matrix-header {
            background: linear-gradient(135deg, #155799, #159957);
            color: white !important;
            padding: 15px 20px;
            border-radius: 8px 8px 0 0;
            margin: -25px -25px 20px -25px;
        }
        .matrix-header h2 {
            margin: 0;
            font-weight: 400;
        }
        .matrix-table {
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 25px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.07);
        }
        .matrix-table th {
            background-color: #f8f9fa;
            color: #333;
            padding: 12px 15px;
            font-weight: 600;
            text-align: left;
            border: 1px solid #e9ecef;
            position: sticky;
            top: 0;
            z-index: 10;
        }
        .matrix-table td {
            border: 1px solid #e9ecef;
            padding: 12px 15px;
            vertical-align: top;
        }
        .matrix-table .control-id {
            font-weight: 600;
            font-family: 'Courier New', monospace;
        }
        .matrix-table .control-desc {
            color: #555;
            font-size: 0.9em;
        }
        .status-dot {
            display: inline-block;
            height: 20px;
            width: 20px;
            border-radius: 50%;
            margin-right: 10px;
            vertical-align: middle;
        }
        .status-label {
            display: inline-block;
            vertical-align: middle;
            font-weight: 600;
        }
        .matrix-type-header {
            background: linear-gradient(135deg, #3a6186, #4568DC);
            color: white !important;
            padding: 12px 15px;
            font-size: 1.2em;
            border-radius: 6px;
            margin: 30px 0 15px 0;
        }
        .matrix-legend {
            display: flex;
            justify-content: flex-end;
            margin-bottom: 15px;
            gap: 15px;
        }
        .legend-item {
            display: flex;
            align-items: center;
            font-size: 0.9em;
        }
        .tooltip {
            position: relative;
            display: inline-block;
            cursor: pointer;
        }
        .tooltip .tooltiptext {
            visibility: hidden;
            width: 300px;
            background-color: #333;
            color: #fff;
            text-align: left;
            border-radius: 6px;
            padding: 10px;
            position: absolute;
            z-index: 1;
            bottom: 125%;
            left: 50%;
            transform: translateX(-50%);
            opacity: 0;
            transition: opacity 0.3s;
            box-shadow: 0 4px 8px rgba(0,0,0,0.2);
        }
        .tooltip:hover .tooltiptext {
            visibility: visible;
            opacity: 1;
        }
        </style>
        
        <div class="matrix-container">
            <div class="matrix-header">
                <h2>UAE IA Controls Compliance Matrix</h2>
            </div>
            
            <div class="matrix-legend">
                <div class="legend-item">
                    <span class="status-dot" style="background-color: #27ae60;"></span>
                    <span>Fully Addressed</span>
                </div>
                <div class="legend-item">
                    <span class="status-dot" style="background-color: #f39c12;"></span>
                    <span>Partially Addressed</span>
                </div>
                <div class="legend-item">
                    <span class="status-dot" style="background-color: #e74c3c;"></span>
                    <span>Not Addressed</span>
                </div>
            </div>
        """
        
        # For each control type, create a section and table
        for control_type, controls in controls_by_type.items():
            if not controls:
                continue
                
            html += f'<div class="matrix-type-header">{control_type}</div>'
            html += '<table class="matrix-table">'
            html += '<tr><th>Control ID</th><th>Description</th>'
            for doc in document_names:
                html += f'<th>{doc}</th>'
            html += '</tr>'
            
            # Control rows
            for control_data in sorted(controls, key=lambda x: x['control']['id']):
                control = control_data['control']
                html += '<tr>'
                html += f'<td class="control-id">{control["id"]}</td>'
                
                desc = control.get('description', '')
                if len(desc) > 100:
                    desc = desc[:100] + '...'
                html += f'<td class="control-desc">{desc}</td>'
                
                # Status for each document
                for doc in document_names:
                    status_data = control_data['status'].get(doc, {
                        'status': 'Not Analyzed',
                        'color': '#a9a9a9',
                        'reason': '',
                        'evidence': ''
                    })
                    
                    status = status_data.get('status', 'Not Analyzed')
                    color = status_data.get('color', '#a9a9a9')
                    reason = status_data.get('reason', '')
                    evidence = status_data.get('evidence', '')
                    
                    # Create tooltip with details if available
                    if reason or evidence:
                        tooltip_content = f"<strong>Status:</strong> {status}<br>"
                        if reason:
                            tooltip_content += f"<strong>Reason:</strong> {reason}<br>"
                        if evidence:
                            tooltip_content += f"<strong>Evidence:</strong> {evidence}"
                            
                        html += f"""
                        <td>
                            <div class="tooltip">
                                <span class="status-dot" style="background-color: {color};"></span>
                                <span class="tooltiptext">{tooltip_content}</span>
                            </div>
                        </td>
                        """
                    else:
                        html += f'<td><span class="status-dot" style="background-color: {color};"></span></td>'
                
                html += '</tr>'
            
            html += '</table>'
        
        html += '</div>'
        
        return html
    
    except Exception as e:
        print(f"Error creating matrix report: {str(e)}")
        traceback.print_exc()
        return f"<p>Error generating matrix report: {str(e)}</p>"

def generate_enhanced_summary(all_results):
    """Generate an enhanced summary dashboard with better visualization."""
    if not all_results or len(all_results) == 0:
        return "<p>No data available for summary dashboard</p>"
    
    try:
        # Collect overall statistics
        total_implemented = []
        total_missing = []
        total_partially_addressed = []
        
        for result in all_results:
            total_implemented.extend(result.get("implemented", []))
            total_partially_addressed.extend(result.get("partially_addressed", []))
            total_missing.extend(result.get("missing", []))
        
        total_all = len(total_implemented) + len(total_missing) + len(total_partially_addressed)
        
        # Calculate percentages
        if total_all > 0:
            impl_pct = len(total_implemented) / total_all * 100
            partial_pct = len(total_partially_addressed) / total_all * 100
            miss_pct = len(total_missing) / total_all * 100
        else:
            impl_pct = partial_pct = miss_pct = 0
        
        # Group controls by type
        controls_by_type = {}
        # Include all controls (implemented, partially addressed, and missing)
        for control in total_implemented + total_partially_addressed + total_missing:
            control_type = control.get("type", "Unknown")
            if control_type not in controls_by_type:
                controls_by_type[control_type] = {
                    "total": 0,
                    "implemented": 0,
                    "partially_addressed": 0,
                    "not_implemented": 0
                }
        
        # Count by implementation status for each type
        for control in total_implemented:
            control_type = control.get("type", "Unknown")
            controls_by_type[control_type]["total"] += 1
            controls_by_type[control_type]["implemented"] += 1
        
        for control in total_partially_addressed:
            control_type = control.get("type", "Unknown")
            controls_by_type[control_type]["total"] += 1
            controls_by_type[control_type]["partially_addressed"] += 1
            
        for control in total_missing:
            control_type = control.get("type", "Unknown")
            controls_by_type[control_type]["total"] += 1
            controls_by_type[control_type]["not_implemented"] += 1
        
        # Create HTML for the dashboard - using f-string formatting
        html = f"""
        <style>
            .dashboard-container {{
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                background-color: #f8f9fa;
                border-radius: 8px;
                overflow: hidden;
                box-shadow: 0 4px 15px rgba(0,0,0,0.1);
            }}
            
            .dashboard-header {{
                background: linear-gradient(135deg, #1a5276, #2471a3);
                color: white !important;
                padding: 20px 25px;
            }}
            
            .dashboard-header h2 {{
                margin: 0;
                color: white !important;
                font-weight: 400;
                font-size: 1.8em;
            }}
            
            .dashboard-header p {{
                margin: 10px 0 0 0;
                color: white !important;
                opacity: 0.9;
            }}
            
            .dashboard-content {{
                padding: 25px;
                color:black;
                background-color: white;
            }}
            
            .stats-overview {{
                display: flex;
                flex-wrap: wrap;
                gap: 20px;
                margin-bottom: 30px;
            }}
            
            .stat-card {{
                flex: 1;
                min-width: 200px;
                background-color: white;
                border-radius: 8px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.07);
                padding: 20px;
                text-align: center;
                transition: transform 0.2s ease;
            }}
            
            .stat-card:hover {{
                transform: translateY(-5px);
                box-shadow: 0 5px 15px rgba(0,0,0,0.1);
            }}
            
            .stat-value {{
                font-size: 2.5em;
                font-weight: 700;
                margin: 10px 0;
            }}
            
            .stat-label {{
                font-size: 1em;
                color: #555;
                margin-bottom: 10px;
            }}
            
            .status-dot {{
                display: inline-block;
                height: 12px;
                width: 12px;
                border-radius: 50%;
                margin-right: 8px;
            }}
            
            .section-title {{
                margin: 30px 0 15px 0;
                padding-bottom: 10px;
                border-bottom: 1px solid #e9ecef;
                color: #2c3e50;
                font-weight: 500;
            }}
            
            .type-analysis {{
                display: flex;
                flex-wrap: wrap;
                gap: 15px;
                margin-top: 30px;
            }}
            
            .progress-container {{
                width: 100%;
                background-color: #444;
                border-radius: 10px;
                height: 10px;
                margin: 10px 0;
            }}
            
            .progress-bar {{
                height: 100%;
                border-radius: 10px;
            }}
            
            .type-card {{
                flex: 1;
                min-width: 250px;
                border-radius: 8px;
                overflow: hidden;
                box-shadow: 0 2px 10px rgba(0,0,0,0.07);
                margin-bottom: 20px;
                margin-right: 15px;
            }}
            
            .type-header {{
                background: linear-gradient(135deg, #3a6186, #4568DC);
                color: white !important;
                padding: 15px;
            }}
            
            .type-content {{
                padding: 15px;
                color: black !important;
                background-color: white;
            }}
            
            .highlight {{
                font-weight: 600;
                font-size: 1.1em;
            }}
            
            .chart-container {{
                width: 100%;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                background-color: white;
                border-radius: 8px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.07);
            }}
            
            .pie-chart {{
                width: 200px;
                height: 200px;
                border-radius: 50%;
                background: conic-gradient(
                    #27ae60 0% {impl_pct}%,
                    #f39c12 {impl_pct}% {impl_pct + partial_pct}%,
                    #e74c3c {impl_pct + partial_pct}% 100%
                );
                margin: 20px auto;
                position: relative;
            }}
            
            .pie-center {{
                position: absolute;
                width: 120px;
                height: 120px;
                background: white;
                border-radius: 50%;
                top: 50%;
                left: 50%;
                transform: translate(-50%, -50%);
                display: flex;
                align-items: center;
                justify-content: center;
                font-weight: 600;
                color: #2c3e50;
            }}
            
            .chart-legend {{
                display: flex;
                justify-content: center;
                flex-wrap: wrap;
                gap: 20px;
                margin-top: 20px;
            }}
            
            .legend-item {{
                display: flex;
                align-items: center;
                font-size: 0.9em;
            }}
            
            .document-row {{
                transition: background-color 0.2s ease;
            }}
            
            .document-row:hover {{
                background-color: #f5f9ff;
            }}
            
            .document-details {{
                display: flex;
                align-items: center;
            }}
            
            .doc-icon {{
                margin-right: 10px;
                color: #3498db;
            }}
            
            .doc-progress {{
                height: 8px;
                background: #eee;
                width: 100%;
                border-radius: 4px;
                margin-top: 6px;
                overflow: hidden;
                position: relative;
            }}
            
            .doc-progress-bar {{
                position: absolute;
                height: 100%;
                background: linear-gradient(to right, #27ae60, #f39c12, #e74c3c);
                border-radius: 4px;
            }}
        </style>
        
        <div class="dashboard-container">
            <div class="dashboard-header">
                <h2>UAE IA Compliance Dashboard</h2>
                <p>Summary of document analysis against compliance controls</p>
            </div>
            
            <div class="dashboard-content">
                <!-- Overview Stats -->
                <div class="stats-overview">
                    <div class="stat-card">
                        <div class="stat-label">Total Controls Checked</div>
                        <div class="stat-value">{total_all}</div>
                    </div>
                    
                    <div class="stat-card">
                        <div class="stat-label">Fully Addressed</div>
                        <div class="stat-value" style="color: #27ae60;">{len(total_implemented)}</div>
                    </div>
                    
                    <div class="stat-card">
                        <div class="stat-label">Partially Addressed</div>
                        <div class="stat-value" style="color: #f39c12;">{len(total_partially_addressed)}</div>
                    </div>
                    
                    <div class="stat-card">
                        <div class="stat-label">Not Addressed</div>
                        <div class="stat-value" style="color: #e74c3c;">{len(total_missing)}</div>
                    </div>
                </div>
                
                <!-- Compliance Chart -->
                <h3 class="section-title">Overall Compliance Status</h3>
                <div class="chart-container">
                    <div class="pie-chart">
                        <div class="pie-center">{total_all}<br>Controls</div>
                    </div>
                    <div class="chart-legend">
                        <div class="legend-item">
                            <span class="status-dot" style="background-color: #27ae60;"></span>
                            <span>Fully Addressed ({impl_pct:.1f}%)</span>
                        </div>
                        <div class="legend-item">
                            <span class="status-dot" style="background-color: #f39c12;"></span>
                            <span>Partially Addressed ({partial_pct:.1f}%)</span>
                        </div>
                        <div class="legend-item">
                            <span class="status-dot" style="background-color: #e74c3c;"></span>
                            <span>Not Addressed ({miss_pct:.1f}%)</span>
                        </div>
                    </div>
                </div>
                
                <!-- Control Type Analysis -->
                <h3 class="section-title">Control Type Analysis</h3>
                <div class="type-analysis">
        """
        
        # Add cards for each control type
        for control_type, stats in controls_by_type.items():
            total = stats["total"]
            implemented = stats["implemented"]
            partially_addressed = stats["partially_addressed"]
            not_addressed = stats["not_implemented"]
            
            if total > 0:
                impl_pct_type = (implemented / total) * 100
                partial_pct_type = (partially_addressed / total) * 100
            else:
                impl_pct_type = 0
                partial_pct_type = 0
                
            html += f"""
                <div class="type-card">
                    <div class="type-header">
                        <h4 style="margin: 0;">{control_type}</h4>
                    </div>
                    <div class="type-content">
                        <p>Total Controls: <span class="highlight">{total}</span></p>
                        <div class="progress-container">
                            <div class="progress-bar" style="width: {(implemented + partially_addressed)/total*100 if total > 0 else 0}%; background-color: #27ae60;"></div>
                        </div>
                        <div style="display: flex; justify-content: space-between; margin-top: 15px;">
                            <div>
                                <div style="color: #27ae60; font-weight: 600;">
                                    <span class="progress-bar" style="display: inline-block; width: 20px; height: 10px; background-color: #27ae60;"></span> {implemented}
                                </div>
                                <div style="font-size: 0.8em; color: #555;">Fully Addressed</div>
                            </div>
                            <div>
                                <div style="color: #f39c12; font-weight: 600;">
                                    <span class="progress-bar" style="display: inline-block; width: 20px; height: 10px; background-color: #f39c12;"></span> {partially_addressed}
                                </div>
                                <div style="font-size: 0.8em; color: #555;">Partially</div>
                            </div>
                            <div>
                                <div style="color: #e74c3c; font-weight: 600;">
                                    <span class="progress-bar" style="display: inline-block; width: 20px; height: 10px; background-color: #e74c3c;"></span> {not_addressed}
                                </div>
                                <div style="font-size: 0.8em; color: #555;">Not Addressed</div>
                            </div>
                        </div>
                    </div>
                </div>
                """
        
        html += """
                </div>
            </div>
        </div>
        """
        
        return html
        
    except Exception as e:
        print(f"Error creating enhanced summary: {str(e)}")
        traceback.print_exc()
        return f"<p>Error generating enhanced summary: {str(e)}</p>"


def process_with_progress(files, progress=gr.Progress()):
    """Process files with enhanced progress reporting."""
    if files is None or len(files) == 0:
        return (
            "<p>Please upload at least one document file.</p>", 
            "<p>No files selected.</p>", 
            "<p>No files selected.</p>",
            [], [], [], "", []
        )
    
    # Create a wrapper for the progress function
    progress_wrapper = ProgressWrapper()
    
    start_time = time.time()
    print(f"Processing {len(files)} files: {[os.path.basename(f) for f in files]}")
    
    try:
        # Start with file preparation
        print("Initializing analysis...")
        
        # Run the analysis
        results = process_multiple_uploads(files)
        
        # Generate reports with progress updates
        print("Generating comprehensive reports...")
        combined_report_html = generate_combined_report_html(results[8])
        
        print("Creating controls matrix visualization...")
        matrix_html = generate_matrix_report_html(results[8])
        
        print("Building summary dashboard...")
        enhanced_summary = generate_enhanced_summary(results[8])
        
        # Calculate and log total time
        total_time = time.time() - start_time
        minutes = int(total_time // 60)
        seconds = int(total_time % 60)
        
        print(f"✅ Analysis complete in {minutes}m {seconds}s")
        
        return [
            enhanced_summary,
            matrix_html,
            combined_report_html,
            results[4],
            results[5],
            results[6],
            results[7],
            results[8]
        ]
    except Exception as e:
        print(f"Error in process_with_progress: {str(e)}")
        traceback.print_exc()
        return ["Error processing files: " + str(e)] * 8
    
def create_ui():
    """Create a professional Gradio interface with improved layout."""
    # Preload the model at startup
    load_fine_tuned_model()
    # Preload controls at startup
    load_controls()
    
    with gr.Blocks(
        title="UAE IA Compliance Analyzer",
        css="""
            :root {
                --primary-color: #1a5276;
                --secondary-color: #2471a3;
                --accent-color: #3498db;
                --success-color: #27ae60;
                --warning-color: #f39c12;
                --danger-color: #e74c3c;
                --light-bg: #f5f8fa;
                --dark-text: #2c3e50;
                --border-radius: 8px;
                --box-shadow: 0 4px 12px rgba(0,0,0,0.1);
            }
            
            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                background-color: var(--light-bg);
            }
            
            .container { 
                max-width: 1300px; 
                margin: 0 auto; 
                padding: 0 20px;
            }
            
            /* Custom style for file uploader */
            .file-uploader {
                max-width: 800px !important;
                margin: 0 auto !important;
            }
            
            /* Style the file upload area */
            .upload-area {
                border: 2px dashed #cbd5e0 !important;
                border-radius: var(--border-radius) !important;
                padding: 20px 15px !important;
                text-align: center !important;
                transition: all 0.3s ease !important;
                background-color: #f8fafc !important;
                margin: 0 auto !important;
                /* Limit height */
                max-height: 180px !important;
                overflow-y: auto !important;
            }
            
            /* Make file upload more compact */
            .file-preview {
                margin-bottom: 8px !important;
                padding: 6px !important;
            }
            
            .file-preview-name {
                font-size: 0.9em !important;
            }
            
            /* Improved progress bar */
            .progress-bar {
                height: 12px !important;
                border-radius: 6px !important;
                margin: 10px 0 !important;
                background-color: #e9ecef !important;
                overflow: hidden !important;
            }
            
            .progress-bar-inner {
                background: linear-gradient(90deg, var(--accent-color), var(--primary-color)) !important;
                height: 100% !important;
                border-radius: 6px !important;
                transition: width 0.3s ease !important;
            }
            
            .progress-text {
                font-size: 0.9em !important;
                margin-top: 5px !important;
                color: var(--dark-text) !important;
                font-weight: 500 !important;
            }
            
            .header { 
                background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
                color: white !important; 
                padding: 20px 25px; 
                margin-bottom: 25px; 
                border-radius: var(--border-radius);
                box-shadow: var(--box-shadow);
            }
            
            .header h1 { 
                margin: 0; 
                color: white !important;
                font-weight: 400;
                font-size: 2.2em;
            }
            
            .header p { 
                margin: 10px 0 0 0; 
                opacity: 0.9; 
                color: white !important;
                line-height: 1.4;
                font-size: 1em;
            }
            
            .card { 
                background-color: white; 
                border-radius: var(--border-radius); 
                box-shadow: var(--box-shadow);
                padding: 25px; 
                margin-bottom: 25px; 
                transition: transform 0.2s ease;
            }
            
            .card:hover {
                transform: translateY(-3px);
                box-shadow: 0 6px 15px rgba(0,0,0,0.15);
            }
            
            .analyze-btn { 
                background: linear-gradient(to right, var(--accent-color), var(--secondary-color)) !important;
                border: none !important;
                color: white !important;
                font-weight: 500 !important;
                padding: 12px 25px !important;
                border-radius: 6px !important;
                box-shadow: 0 4px 10px rgba(52, 152, 219, 0.3) !important;
                transition: all 0.3s ease !important;
                width: 600px !important;  /* Match the file uploader width */
                max-width: 100% !important;
                margin: 15px auto !important;
                display: block !important;
                font-size: 1.1em !important;
            }
            
            /* Rest of your CSS */
        """
    ) as app:
        gr.HTML("""
        <div class="header">
            <h1>UAE IA Compliance Analyzer</h1>
            <p>Analyze documents against UAE Information Assurance controls, standards, and policies.</p>
        </div>
        """)
        
        # State variables to store analysis results
        implemented_state = gr.State([])
        missing_state = gr.State([])
        partially_state = gr.State([])
        filename_state = gr.State("")
        all_results_state = gr.State([])
        
        # Create a centered column with file upload area
        with gr.Row(elem_classes=["file-uploader-row"]):
            with gr.Column(elem_classes=["file-uploader"]):
                file_input = gr.File(
                    label="Upload Documents (PDF or DOCX)",
                    file_count="multiple",
                    elem_classes=["upload-area"],
                    # Don't use width/height parameters, use CSS instead
                )
                
                analyze_btn = gr.Button("Analyze Documents", elem_classes=["analyze-btn"])
        
        # Create tabs for different views
        with gr.Tabs() as tabs:
            with gr.Tab("Summary Dashboard"):
                summary_output = gr.HTML(label="Analysis Summary")
            
            with gr.Tab("Matrix View"):
                matrix_view_output = gr.HTML(label="Controls Matrix")
                
            with gr.Tab("Detailed Report"):
                combined_report_output = gr.HTML(label="Combined Report")
        
        # Connect the analysis button
        analyze_btn.click(
            fn=process_with_progress,
            inputs=[file_input],
            outputs=[
                summary_output,
                matrix_view_output,
                combined_report_output,
                implemented_state, 
                missing_state, 
                partially_state,
                filename_state,
                all_results_state
            ],
            show_progress="full"
        )
           
        gr.HTML("""
        <div class="footer">
            <p>UAE IA Compliance Analyzer | Powered by Fine-Tuned llama3.2 Model</p>
        </div>
        """)
    
    return app
if __name__ == "__main__":
    # Enable memory optimization for PyTorch
    if torch.cuda.is_available():
        # Set to reduce memory fragmentation
        torch.cuda.empty_cache()
    
    # Debug file path
    file_path = '/home/gpu/Documents/Amal/Compliance_llama3.2/control/controlt.json'
    print(f"Does file exist? {os.path.exists(file_path)}")
    
    # Create and launch the UI
    app = create_ui()
    app.launch(server_name=os.environ.get('HOST', '0.0.0.0'),
               server_port=int(os.environ.get('PORT', 8880)), root_path="/gradio-demo",
               debug=False, share=False)